import pymysql
from datetime import datetime
import pandas as pd
from docx import Document
from docx.shared import Inches
import io
import traceback
import logging
import sys
import os

# 配置日志
logging.basicConfig(
    level=logging.DEBUG,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(sys.stdout),
        logging.FileHandler('bit_test.log')
    ]
)
logger = logging.getLogger(__name__)

class BitTestService:
    def __init__(self):
        self.db_config = {
            'host': 'localhost',
            'user': 'root',
            'password': '123456',
            'database': 'webproject',
            'charset': 'utf8mb4'
        }
        self.connection = None
        self.is_testing = False
        logger.info("BitTestService initialized")

    def get_connection(self):
        try:
            if not self.connection:
                logger.debug("Creating new database connection")
                self.connection = pymysql.connect(**self.db_config)
            return self.connection
        except Exception as e:
            logger.error(f"Database connection error: {str(e)}")
            raise

    def get_test_list(self, page, size):
        try:
            conn = self.get_connection()
            with conn.cursor(pymysql.cursors.DictCursor) as cursor:
                # 获取总记录数
                cursor.execute("SELECT COUNT(*) as total FROM pms_sca")
                total = cursor.fetchone()['total']

                # 获取分页数据
                offset = (page - 1) * size
                cursor.execute("""
                    SELECT 
                        item as projectName,
                        project_type as type,
                        status as result,
                        info as description
                    FROM pms_sca
                    ORDER BY id DESC
                    LIMIT %s OFFSET %s
                """, (size, offset))
                
                items = cursor.fetchall()
                
                return {
                    'items': items,
                    'total': total
                }
        except Exception as e:
            print(f"Error in get_test_list: {str(e)}")
            raise e

    def start_test(self):
        if self.is_testing:
            raise Exception("测试已经在进行中")
        self.is_testing = True
        # TODO: 实现实际的测试逻辑

    def stop_test(self):
        if not self.is_testing:
            raise Exception("没有正在进行的测试")
        self.is_testing = False
        # TODO: 实现停止测试的逻辑

    def export_excel(self, file_path):
        try:
            logger.info(f"Starting Excel export to {file_path}")
            conn = self.get_connection()
            
            with conn.cursor(pymysql.cursors.DictCursor) as cursor:
                logger.debug("Executing SQL query")
                cursor.execute("""
                    SELECT 
                        item as projectName,
                        project_type as type,
                        status as result,
                        info as description
                    FROM pms_sca
                    ORDER BY id DESC
                """)
                
                data = cursor.fetchall()
                logger.debug(f"Fetched {len(data)} records from database")
                
                if not data:
                    logger.warning("No data found for export")
                    raise Exception("没有数据可导出")
                
                # 创建DataFrame
                logger.debug("Creating DataFrame")
                df = pd.DataFrame(data)
                logger.debug(f"DataFrame created with shape: {df.shape}")
                
                # 确保目录存在
                os.makedirs(os.path.dirname(file_path), exist_ok=True)
                
                # 直接保存到文件
                logger.debug(f"Saving Excel file to {file_path}")
                with pd.ExcelWriter(file_path, engine='openpyxl') as writer:
                    df.to_excel(writer, index=False, sheet_name='符合性测试报告')
                    
                    # 自动调整列宽
                    logger.debug("Adjusting column widths")
                    worksheet = writer.sheets['符合性测试报告']
                    for idx, col in enumerate(df.columns):
                        max_length = max(
                            df[col].astype(str).apply(len).max(),
                            len(col)
                        )
                        worksheet.column_dimensions[chr(65 + idx)].width = max_length + 2
                
                logger.info("Excel file saved successfully")
                return True
                
        except Exception as e:
            logger.error(f"导出Excel失败: {str(e)}")
            logger.error(f"错误详情: {traceback.format_exc()}")
            raise Exception(f"导出Excel失败: {str(e)}")

    def export_word(self):
        try:
            logger.info("Starting Word export")
            conn = self.get_connection()
            
            with conn.cursor(pymysql.cursors.DictCursor) as cursor:
                logger.debug("Executing SQL query")
                cursor.execute("""
                    SELECT 
                        item,
                        project_type,
                        status,
                        info
                    FROM pms_sca
                    ORDER BY id DESC
                """)
                
                data = cursor.fetchall()
                logger.debug(f"Fetched {len(data)} records from database")
                
                if not data:
                    logger.warning("No data found for export")
                    raise Exception("没有数据可导出")
                
                # 创建Word文档
                logger.debug("Creating Word document")
                doc = Document()
                doc.add_heading('符合性测试报告', 0)
                doc.add_paragraph(f'生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
                
                # 添加表格
                logger.debug("Adding table to Word document")
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Table Grid'
                
                # 添加表头
                header_cells = table.rows[0].cells
                header_cells[0].text = '项目名称'
                header_cells[1].text = '类型'
                header_cells[2].text = '结果'
                header_cells[3].text = '描述'
                
                # 添加数据行
                logger.debug("Adding data rows to table")
                for row in data:
                    cells = table.add_row().cells
                    cells[0].text = str(row['item'])
                    cells[1].text = str(row['project_type'])
                    cells[2].text = str(row['status'])
                    cells[3].text = str(row['info'])
                
                # 保存到内存
                logger.debug("Saving Word document to memory")
                output = io.BytesIO()
                doc.save(output)
                output.seek(0)
                word_data = output.getvalue()
                logger.debug(f"Word data size: {len(word_data)} bytes")
                
                logger.info("Word export completed successfully")
                return word_data
                
        except Exception as e:
            logger.error(f"导出Word失败: {str(e)}")
            logger.error(f"错误详情: {traceback.format_exc()}")
            raise Exception(f"导出Word失败: {str(e)}")

    def __del__(self):
        if self.connection:
            logger.debug("Closing database connection")
            self.connection.close() 